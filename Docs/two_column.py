from docx import Document
from docx.shared import Inches

def split_text_by_separator(text, separator="==="):
    """Разделяет текст по заданному разделителю"""
    parts = text.split(separator)
    if len(parts) == 2:
        return parts[0].strip(), parts[1].strip()
    else:
        raise ValueError("Не найдено ровно два блока текста. Проверьте разделитель.")

def read_docx(file_path):
    """Читает текст из .docx файла и возвращает его в виде строки"""
    doc = Document(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def create_two_column_docx(english_text, translated_text, output_path):
    """Создает документ с двумя колонками"""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(3)  # Левая колонка
    table.columns[1].width = Inches(3)  # Правая колонка

    row = table.rows[0]
    row.cells[0].text = english_text
    row.cells[1].text = translated_text

    doc.save(output_path)
    print(f"Документ сохранен как {output_path}")

def process_docx(input_path, output_path, separator="==="):
    """Основная функция обработки файла"""
    text = read_docx(input_path)
    english_text, translated_text = split_text_by_separator(text, separator)
    create_two_column_docx(english_text, translated_text, output_path)

# Пример использования
input_file = "input.docx"   # Исходный файл
output_file = "output.docx" # Файл с двумя колонками
process_docx(input_file, output_file)
