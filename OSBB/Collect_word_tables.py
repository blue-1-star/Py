import os
from docx import Document
from openpyxl import Workbook
from openpyxl.utils import get_column_letter


def batch_word_tables_to_excel(word_dir, output_excel_path):
    """Проходит по всем файлам .

    docx в каталоге, извлекает первую таблицу из каждого и сохраняет на
    отдельный лист в общий файл Excel.
    """
    if not os.path.exists(word_dir):
        print(f"Ошибка: Указанный каталог {word_dir} не найден.")
        return

    # 1. Создаем новую рабочую книгу Excel
    wb = Workbook()
    # Удаляем пустой дефолтный лист, создаваемый автоматически
    default_sheet = wb.active
    wb.remove(default_sheet)

    # 2. Получаем список всех файлов .docx в папке
    # Фильтруем только .docx и отсекаем временные файлы Windows (начинающиеся с ~$)
    word_files = [
        f
        for f in os.listdir(word_dir)
        if f.endswith(".docx") and not f.startswith("~$")
    ]

    if not word_files:
        print(f"В каталоге {word_dir} не найдено файлов .docx для обработки.")
        return

    print(f"Найдено файлов Word для обработки: {len(word_files)}")

    # 3. Цикл по всем найденным файлам
    for file_name in word_files:
        word_path = os.path.join(word_dir, file_name)

        try:
            doc = Document(word_path)

            # Проверяем, есть ли в документе таблицы
            if not doc.tables:
                print(
                    f"Предупреждение: В файле '{file_name}' нет таблиц. Пропускаем."
                )
                continue

            # По условию таблица одна, берем самую первую (индекс 0)
            table = doc.tables[0]

            # Формируем имя листа: убираем ".docx" из названия файла
            sheet_title = os.path.splitext(file_name)[0]

            # Очищаем имя от запрещенных в Excel символов
            for char in [":", "?", "*", "/", "\\"]:
                sheet_title = sheet_title.replace(char, "")

            # Excel строго ограничивает длину имени листа в 31 символ
            sheet_title = sheet_title[:31]

            # Создаем новый лист с именем файла
            ws = wb.create_sheet(title=sheet_title)
            print(f"Обработка: '{file_name}' -> Лист Excel: '{sheet_title}'")

            # Читаем строки таблицы из Word и записываем в Excel
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                ws.append(row_data)

            # Автоматическое выравнивание ширины колонок под текст
            for col in ws.columns:
                max_len = max(len(str(cell.value or "")) for cell in col)
                col_letter = get_column_letter(col[0].column)
                ws.column_dimensions[col_letter].width = max(max_len + 3, 10)

        except Exception as e:
            print(f"Ошибка при обработке файла {file_name}: {e}")

    # 4. Проверяем, записался ли хоть один лист, перед сохранением итогового файла
    if not wb.sheetnames:
        print(
            "Ни одной таблицы не удалось импортировать. Файл Excel создан не будет."
        )
        return

    # 5. Сохраняем объединенную книгу Excel
    wb.save(output_excel_path)
    print(f"\nУспех! Все данные сохранены в один файл: {output_excel_path}")


# --- Настройки путей ---
if __name__ == "__main__":
    # Папка, где лежат ваши файлы «5 під'їзд.docx» и другие
    # Буква 'r' перед кавычками обязательна, чтобы слэши не ломали путь!
    input_directory = r"G:\Programming\Py\OSBB\Data\raw\typed"

    # Имя итогового файла Excel (сохранится в той же папке)
    final_excel = r"G:\Programming\Py\OSBB\Data\raw\typed\p_all_tables.xlsx"

    # Запуск функции
    batch_word_tables_to_excel(input_directory, final_excel)
