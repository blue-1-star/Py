import os
from docx import Document
from openpyxl import Workbook
from openpyxl.utils import get_column_letter


def word_tables_to_excel(word_path, excel_path):
    # 1. Загружаем документ Word
    if not os.path.exists(word_path):
        print(f"Ошибка: Файл {word_path} не найден.")
        return

    doc = Document(word_path)

    if not doc.tables:
        print("В документе Word не найдено ни одной таблицы.")
        return

    # 2. Создаем новую рабочую книгу Excel
    wb = Workbook()
    # Удаляем дефолтный лист, который создается автоматически
    default_sheet = wb.active
    wb.remove(default_sheet)

    print(f"Найдено таблиц в Word: {len(doc.tables)}")

    # 3. Перебираем таблицы и переносим данные
    for index, table in enumerate(doc.tables, start=1):
        # Создаем новый лист для каждой таблицы
        sheet_title = f"Таблица {index}"
        ws = wb.create_sheet(title=sheet_title)

        # Читаем строки и ячейки из таблицы Word
        for row in table.rows:
            # Собираем текст из каждой ячейки строки
            row_data = [cell.text.strip() for cell in row.cells]
            ws.append(row_data)

        # Автоматическое выравнивание ширины колонок для красоты
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            col_letter = get_column_letter(col[0].column)
            # Задаем ширину с небольшим запасом
            ws.column_dimensions[col_letter].width = max(max_len + 3, 10)

        print(f"Добавлен лист: {sheet_title}")

    # 4. Сохраняем итоговый файл Excel
    wb.save(excel_path)
    print(f"Готово! Данные успешно сохранены в файл: {excel_path}")


# --- Пример использования ---
if __name__ == "__main__":
    # Укажите пути к вашим файлам
    # word_file = "document.docx"  # Имя вашего файла Word
    word_file =r"G:\Programming\Py\OSBB\Data\raw\typed\5 під'їзд.docx"
    # excel_file = "extracted_tables.xlsx"  # Имя итогового файла Excel
    excel_file = r"G:\Programming\Py\OSBB\Data\raw\typed\p_5.xlsx"

    word_tables_to_excel(word_file, excel_file)
